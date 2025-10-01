"""
Professional Report Generator - FIXED VERSION
Robust PDF, HTML, and DOCX reports with comprehensive content display
"""

import pandas as pd
import numpy as np
from typing import Dict, List, Any, Optional
from pathlib import Path
import base64
import io
from datetime import datetime
from loguru import logger
from docx.shared import RGBColor
import warnings
import tempfile
import os
import sys

warnings.filterwarnings('ignore')

# Import handling with robust error prevention
try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("FPDF not available for PDF generation")

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available for DOCX generation")

try:
    import matplotlib.pyplot as plt
    import seaborn as sns
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

try:
    from jinja2 import Template
    JINJA_AVAILABLE = True
except ImportError:
    JINJA_AVAILABLE = False

class ProfessionalPDF(FPDF):
    """Professional PDF generator with reliable formatting"""
    
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)
        
    def header(self):
        """Professional header"""
        if self.page_no() == 1:
            return
        self.set_font('Arial', 'B', 10)
        self.set_text_color(0, 0, 0)
        self.cell(0, 10, 'Comprehensive Data Analysis Report', 0, 1, 'C')
    
    def footer(self):
        """Professional footer"""
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

class ReportGenerator:
    """Professional report generator with reliable file handling"""
    
    def __init__(self, config=None):
        self.config = config or {}
        self.output_dir = Path("outputs/reports")
        self.temp_dir = Path("temp/charts")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info("Professional Report Generator initialized")
    
    def generate_comprehensive_report(self, 
                                    data: Dict[str, Any],
                                    report_title: str = "Comprehensive Data Analysis Report",
                                    format: str = "html",
                                    include_sections: Dict[str, bool] = None,
                                    include_visualizations: bool = True) -> str:
        """
        Generate comprehensive report with reliable file handling
        """
        try:
            logger.info(f"Generating {format.upper()} report: {report_title}")
            
            # Default sections
            if include_sections is None:
                include_sections = self._get_default_sections()
            
            # Generate report content
            report_content = self._generate_report_content(
                data, report_title, include_sections, include_visualizations
            )
            
            # Generate based on format with robust error handling
            if format.lower() == 'html':
                report_path = self._generate_html_report(report_content, report_title)
            elif format.lower() == 'pdf':
                report_path = self._generate_pdf_report(report_content, report_title, include_visualizations)
            elif format.lower() == 'docx':
                report_path = self._generate_docx_report(report_content, report_title, include_visualizations)
            else:
                report_path = self._generate_html_report(report_content, report_title)
            
            # Verify file was created successfully
            if report_path and self._verify_file_integrity(report_path, format):
                file_size = os.path.getsize(report_path)
                logger.info(f"✅ Report generated successfully: {report_path} ({file_size} bytes)")
                return str(report_path)
            else:
                logger.error("❌ Report generation failed integrity check")
                return ""
                
        except Exception as e:
            logger.error(f"❌ Error generating report: {str(e)}")
            return ""
    
    def _generate_report_content(self, 
                               data: Dict[str, Any], 
                               title: str,
                               include_sections: Dict[str, bool],
                               include_visualizations: bool = True) -> Dict[str, Any]:
        """Generate comprehensive report content"""
        
        content = {
            'title': title,
            'generation_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'sections': {}
        }
        
        # Generate visualizations first
        if include_visualizations and MATPLOTLIB_AVAILABLE:
            content['visualizations'] = self._generate_visualizations(data)
        
        # Generate all sections with proper content - FIXED: Ensure all sections have content
        section_creators = {
            'executive_summary': (self._create_executive_summary, include_sections.get('executive_summary', True)),
            'data_overview': (self._create_data_overview, include_sections.get('data_overview', True)),
            'eda_analysis': (self._create_eda_section, include_sections.get('eda_analysis', True)),
            'ml_results': (self._create_ml_section, include_sections.get('ml_results', True)),
            'explanations': (self._create_explanations_section, include_sections.get('explanations', True)),
            'sql_insights': (self._create_sql_section, include_sections.get('sql_insights', True)),
            'recommendations': (self._create_recommendations_section, include_sections.get('recommendations', True))
        }
        
        for section_key, (creator, should_include) in section_creators.items():
            if should_include:
                try:
                    section_content = creator(data)
                    # Ensure section has content before adding
                    if section_content and section_content.get('content'):
                        content['sections'][section_key] = section_content
                    else:
                        logger.warning(f"Section {section_key} has no content, using fallback")
                        content['sections'][section_key] = self._create_fallback_section(section_key)
                except Exception as e:
                    logger.error(f"Error creating section {section_key}: {str(e)}")
                    content['sections'][section_key] = self._create_error_section(section_key, str(e))
        
        # Ensure we have at least one section
        if not content['sections']:
            content['sections']['fallback'] = self._create_fallback_content()
        
        return content
    
    def _generate_visualizations(self, data: Dict[str, Any]) -> Dict[str, str]:
        """Generate and save visualizations as image files"""
        visualizations = {}
        
        try:
            if 'dataset' in data and MATPLOTLIB_AVAILABLE:
                df = data['dataset']
                
                # Basic dataset info visualization
                plt.figure(figsize=(10, 6))
                
                # Create a simple bar chart showing data types distribution
                if len(df.columns) > 0:
                    dtype_counts = df.dtypes.value_counts()
                    dtype_names = [str(dtype) for dtype in dtype_counts.index]
                    
                    plt.bar(dtype_names, dtype_counts.values, color=['#007acc', '#ff6b6b', '#51cf66', '#fcc419'])
                    plt.title('Data Types Distribution')
                    plt.xlabel('Data Types')
                    plt.ylabel('Count')
                    plt.xticks(rotation=45)
                    plt.tight_layout()
                    
                    dtype_path = self.temp_dir / 'data_types_distribution.png'
                    plt.savefig(dtype_path, dpi=150, bbox_inches='tight', facecolor='white')
                    plt.close()
                    
                    visualizations['data_types'] = str(dtype_path)
                
                # Numeric analysis if available
                numeric_df = df.select_dtypes(include=[np.number])
                if not numeric_df.empty:
                    # Correlation heatmap
                    if len(numeric_df.columns) > 1:
                        plt.figure(figsize=(10, 8))
                        corr_matrix = numeric_df.corr()
                        
                        mask = np.triu(np.ones_like(corr_matrix, dtype=bool))
                        sns.heatmap(corr_matrix, mask=mask, annot=True, cmap='coolwarm', center=0,
                                   square=True, fmt='.2f', cbar_kws={'shrink': 0.8})
                        plt.title('Feature Correlation Matrix', fontsize=14, pad=20)
                        plt.xticks(rotation=45, ha='right')
                        plt.yticks(rotation=0)
                        plt.tight_layout()
                        
                        heatmap_path = self.temp_dir / 'correlation_matrix.png'
                        plt.savefig(heatmap_path, dpi=150, bbox_inches='tight', facecolor='white')
                        plt.close()
                        
                        visualizations['correlation_matrix'] = str(heatmap_path)
            
            # ML visualizations if available
            if 'ml_results' in data and MATPLOTLIB_AVAILABLE:
                ml_data = data['ml_results']
                
                # Create a simple performance chart
                if 'all_models' in ml_data:
                    models = list(ml_data['all_models'].keys())
                    performances = []
                    
                    for model_name, results in ml_data['all_models'].items():
                        metrics = results.get('test_metrics', {})
                        if 'accuracy' in metrics:
                            performances.append(metrics['accuracy'])
                        elif 'r2' in metrics:
                            performances.append(metrics['r2'])
                        else:
                            performances.append(0.0)
                    
                    if performances:
                        plt.figure(figsize=(10, 6))
                        plt.bar(models, performances, color='#007acc', alpha=0.7)
                        plt.title('Model Performance Comparison')
                        plt.xlabel('Models')
                        plt.ylabel('Performance Score')
                        plt.xticks(rotation=45)
                        plt.tight_layout()
                        
                        performance_path = self.temp_dir / 'model_performance.png'
                        plt.savefig(performance_path, dpi=150, bbox_inches='tight', facecolor='white')
                        plt.close()
                        
                        visualizations['model_performance'] = str(performance_path)
            
        except Exception as e:
            logger.error(f"Error generating visualizations: {str(e)}")
            # Don't fail completely if visualizations fail
        
        return visualizations

    def _generate_pdf_report(self, content: Dict[str, Any], title: str, include_visualizations: bool = True) -> Path:
        """Generate reliable PDF report - FIXED VERSION"""
        if not PDF_AVAILABLE:
            logger.warning("PDF generation not available")
            return None
        
        pdf = None
        try:
            pdf = ProfessionalPDF()
            pdf.add_page()
            
            # Title - FIXED: Use simple, reliable formatting
            pdf.set_font('Arial', 'B', 16)
            pdf.cell(0, 10, title, 0, 1, 'C')
            pdf.set_font('Arial', '', 10)
            pdf.cell(0, 10, f"Generated on {content['generation_date']}", 0, 1, 'C')
            pdf.ln(10)
            
            # Main content sections - FIXED: Handle empty sections gracefully
            if not content.get('sections'):
                pdf.multi_cell(0, 10, "No report content available. Please check your data and try again.")
                pdf.ln(10)
            else:
                for section_key, section in content['sections'].items():
                    # FIXED: Ensure section has content
                    if not section or not section.get('content'):
                        continue
                        
                    pdf.set_font('Arial', 'B', 14)
                    pdf.cell(0, 10, section['title'], 0, 1)
                    pdf.set_font('Arial', '', 10)
                    
                    for item in section.get('content', []):
                        try:
                            if item['type'] == 'paragraph':
                                text = self._clean_text(item.get('text', ''))
                                if text:
                                    pdf.multi_cell(0, 6, text)
                                    pdf.ln(2)
                            elif item['type'] == 'bullet':
                                text = self._clean_text(item.get('text', ''))
                                if text:
                                    pdf.cell(5, 6, '', 0, 0)
                                    pdf.cell(5, 6, '•', 0, 0)
                                    pdf.multi_cell(0, 6, text)
                            elif item['type'] == 'numbered':
                                text = self._clean_text(item.get('text', ''))
                                if text:
                                    pdf.cell(5, 6, '', 0, 0)
                                    pdf.cell(10, 6, f"{item.get('number', '')}.", 0, 0)
                                    pdf.multi_cell(0, 6, text)
                            elif item['type'] == 'table':
                                table_title = item.get('title', 'Table')
                                table_data = item.get('data', {})
                                
                                if table_data:
                                    pdf.set_font('Arial', 'B', 11)
                                    pdf.cell(0, 8, table_title, 0, 1)
                                    pdf.set_font('Arial', '', 9)
                                    
                                    for key, value in table_data.items():
                                        pdf.cell(90, 6, str(key), 0, 0)
                                        pdf.cell(90, 6, str(value), 0, 1)
                                    pdf.ln(2)
                        except Exception as item_error:
                            logger.error(f"Error processing PDF item: {str(item_error)}")
                            continue
                    
                    pdf.ln(8)
            
            # Add visualizations if available - FIXED: Handle image paths correctly
            if include_visualizations and content.get('visualizations'):
                pdf.add_page()
                pdf.set_font('Arial', 'B', 14)
                pdf.cell(0, 10, "Data Visualizations", 0, 1)
                pdf.ln(5)
                
                for viz_name, viz_path in content['visualizations'].items():
                    # FIXED: Handle both file paths and base64 data
                    if isinstance(viz_path, str) and viz_path.startswith('data:image'):
                        # Skip base64 images in PDF - they need file paths
                        continue
                    
                    if os.path.exists(viz_path):
                        try:
                            pdf.set_font('Arial', 'B', 10)
                            pdf.cell(0, 8, viz_name.replace('_', ' ').title(), 0, 1)
                            
                            # Safe image embedding with error handling
                            pdf.image(viz_path, x=10, w=190)
                            pdf.ln(5)
                            
                        except Exception as img_error:
                            logger.error(f"PDF image embedding failed for {viz_path}: {str(img_error)}")
                            pdf.multi_cell(0, 6, f"Visualization not available: {viz_name}")
            
            # Generate safe filename
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_title}_{timestamp}.pdf"
            file_path = self.output_dir / filename
            
            # Ensure directory exists
            self.output_dir.mkdir(parents=True, exist_ok=True)
            
            # Save PDF
            pdf.output(str(file_path))
            
            logger.info(f"PDF generated successfully: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"❌ Error generating PDF: {str(e)}")
            if pdf:
                try:
                    pdf.close()
                except:
                    pass
            return None

    def _generate_docx_report(self, content: Dict[str, Any], title: str, include_visualizations: bool = True) -> Path:
        """Generate reliable DOCX report"""
        if not DOCX_AVAILABLE:
            logger.warning("DOCX generation not available")
            return None
        
        try:
            doc = Document()
            
            # Title
            doc.add_heading(title, level=1)
            
            # Date
            doc.add_paragraph(f"Generated on {content['generation_date']}")
            doc.add_paragraph()
            
            # Main content
            for section_key, section in content.get('sections', {}).items():
                if not section or not section.get('content'):
                    continue
                    
                # Add section header with NVIDIA green color
            heading = doc.add_heading(section['title'], level=2)
            
            # Set text color to NVIDIA green (#76B900) for section headers only
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x76, 0xB9, 0x00)  # NVIDIA green

                
                for item in section.get('content', []):
                    try:
                        if item['type'] == 'paragraph':
                            text = self._clean_text(item.get('text', ''))
                            if text:
                                doc.add_paragraph(text)
                        elif item['type'] == 'bullet':
                            text = self._clean_text(item.get('text', ''))
                            if text:
                                p = doc.add_paragraph(text, style='List Bullet')
                        elif item['type'] == 'numbered':
                            text = self._clean_text(item.get('text', ''))
                            if text:
                                p = doc.add_paragraph(text, style='List Number')
                        elif item['type'] == 'table':
                            table_data = item.get('data', {})
                            if table_data:
                                doc.add_heading(item.get('title', 'Table'), level=3)
                                table = doc.add_table(rows=1, cols=2)
                                table.style = 'Table Grid'
                                
                                # Header
                                hdr_cells = table.rows[0].cells
                                hdr_cells[0].text = 'Metric'
                                hdr_cells[1].text = 'Value'
                                
                                # Data rows
                                for key, value in table_data.items():
                                    row_cells = table.add_row().cells
                                    row_cells[0].text = str(key)
                                    row_cells[1].text = str(value)
                    except Exception as item_error:
                        logger.error(f"Error processing DOCX item: {str(item_error)}")
                        continue
                
                doc.add_paragraph()
            
            # Add visualizations
            if include_visualizations and content.get('visualizations'):
                doc.add_heading('Data Visualizations', level=1)
                
                for viz_name, viz_path in content['visualizations'].items():
                    # Handle both file paths and base64 data
                    if isinstance(viz_path, str) and viz_path.startswith('data:image'):
                        continue  # Skip base64 for DOCX
                    
                    if os.path.exists(viz_path):
                        try:
                            doc.add_heading(viz_name.replace('_', ' ').title(), level=2)
                            doc.add_picture(viz_path, width=Inches(6.0))
                            doc.add_paragraph()
                        except Exception as img_error:
                            logger.error(f"DOCX image embedding failed: {str(img_error)}")
            
            # Save file
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_title}_{timestamp}.docx"
            file_path = self.output_dir / filename
            
            self.output_dir.mkdir(parents=True, exist_ok=True)
            doc.save(str(file_path))
            
            return file_path
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            return None

    def _generate_html_report(self, content: Dict[str, Any], title: str) -> Path:
        """Generate comprehensive HTML report with all content visible"""
        if not JINJA_AVAILABLE:
            logger.warning("Jinja2 not available for HTML generation")
            return None
        
        try:
            # Professional HTML template with all content displayed
            html_template = """
            <!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{{ title }}</title>
                <style>
                    body { 
                        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; 
                        margin: 0; 
                        padding: 20px;
                        background-color: #f8f9fa;
                        color: #333;
                        line-height: 1.6;
                    }
                    .container {
                        max-width: 1200px;
                        margin: 0 auto;
                        background: white;
                        padding: 30px;
                        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                        border-radius: 8px;
                    }
                    .header { 
                        text-align: center;
                        margin-bottom: 30px;
                        padding-bottom: 20px;
                        border-bottom: 2px solid #3498db;
                    }
                    .title { 
                        font-size: 2.2em; 
                        margin-bottom: 10px;
                        color: #2c3e50;
                    }
                    .subtitle {
                        font-size: 1.1em;
                        color: #7f8c8d;
                    }
                    .section { 
                        margin: 25px 0; 
                        padding: 20px;
                        background: white;
                        border-radius: 6px;
                        border-left: 4px solid #76B900;
                    }
                    .section-title { 
                        color: #2c3e50; 
                        font-size: 1.4em; 
                        margin-bottom: 15px;
                    }
                    .content-item { 
                        margin: 12px 0; 
                    }
                    .paragraph { 
                        margin: 10px 0; 
                        color: #555;
                    }
                    .bullet { 
                        margin: 8px 0; 
                        padding-left: 20px; 
                        position: relative;
                    }
                    .bullet:before { 
                        content: "•"; 
                        color: #3498db; 
                        font-weight: bold; 
                        position: absolute;
                        left: 0;
                    }
                    .numbered { 
                        margin: 8px 0; 
                        padding-left: 20px; 
                    }
                    .table { 
                        border-collapse: collapse; 
                        width: 100%; 
                        margin: 15px 0; 
                        background: white;
                    }
                    .table th, .table td { 
                        border: 1px solid #e1e1e1; 
                        padding: 10px; 
                        text-align: left; 
                    }
                    .table th { 
                        background: #f8f9fa;
                        color: #2c3e50;
                        font-weight: 600;
                    }
                    .visualization {
                        margin: 20px 0;
                        text-align: center;
                        padding: 15px;
                        background: white;
                        border-radius: 6px;
                        border: 1px solid #e1e1e1;
                    }
                    .visualization img {
                        max-width: 100%;
                        height: auto;
                        border-radius: 4px;
                    }
                    .footer { 
                        text-align: center; 
                        margin-top: 40px; 
                        padding: 20px;
                        color: #7f8c8d;
                        border-top: 1px solid #e1e1e1;
                    }
                </style>
            </head>
            <body>
                <div class="container">
                    <div class="header">
                        <h1 class="title">{{ title }}</h1>
                        <p class="subtitle">Generated on {{ generation_date }}</p>
                    </div>
                    
                    <div class="content-area">
                        {% for section_key, section in sections.items() %}
                        <div class="section">
                            <h2 class="section-title">{{ section.title }}</h2>
                            {% for item in section.content %}
                                <div class="content-item">
                                    {% if item.type == 'paragraph' %}
                                        <p class="paragraph">{{ item.text | safe }}</p>
                                    {% elif item.type == 'bullet' %}
                                        <div class="bullet">{{ item.text | safe }}</div>
                                    {% elif item.type == 'numbered' %}
                                        <div class="numbered">{{ item.number }}. {{ item.text | safe }}</div>
                                    {% elif item.type == 'table' %}
                                        <h4>{{ item.title }}</h4>
                                        <table class="table">
                                            <tbody>
                                                {% for key, value in item.data.items() %}
                                                <tr>
                                                    <th>{{ key }}</th>
                                                    <td>{{ value }}</td>
                                                </tr>
                                                {% endfor %}
                                            </tbody>
                                        </table>
                                    {% endif %}
                                </div>
                            {% endfor %}
                        </div>
                        {% endfor %}
                        
                        {% if visualizations %}
                        <div class="section">
                            <h2 class="section-title">Data Visualizations</h2>
                            {% for viz_name, viz_path in visualizations.items() %}
                            <div class="visualization">
                                <h4>{{ viz_name.replace('_', ' ').title() }}</h4>
                                {% if viz_path.startswith('data:image') %}
                                    <img src="{{ viz_path }}" alt="{{ viz_name }}">
                                {% else %}
                                    <img src="file://{{ viz_path }}" alt="{{ viz_name }}" onerror="this.style.display='none'">
                                {% endif %}
                                <p>{{ viz_name.replace('_', ' ').title() }}</p>
                            </div>
                            {% endfor %}
                        </div>
                        {% endif %}
                    </div>
                    
                    <div class="footer" style="text-align:center; padding: 15px; font-family: 'Segoe UI', sans-serif; font-size: 14px; color: #ccc; background: #0b0c10; border-top: 1px solid #1f2833;">
    <p>Report generated by <strong>GenAI Autonomous Data Agent</strong></p>
    <p>GPU-Accelerated • AI-Powered • Built for <span style="color:#76b900; font-weight:bold;">NVIDIA GTC 2025</span></p>
</div>

                </div>
            </body>
            </html>
            """
            
            template = Template(html_template)
            
            # Convert visualizations to base64 for reliable embedding in HTML
            viz_content = {}
            for viz_name, viz_path in content.get('visualizations', {}).items():
                if os.path.exists(viz_path):
                    try:
                        with open(viz_path, 'rb') as img_file:
                            img_data = base64.b64encode(img_file.read()).decode()
                            viz_content[viz_name] = f"data:image/png;base64,{img_data}"
                    except Exception as e:
                        logger.error(f"Base64 encoding failed: {str(e)}")
                        viz_content[viz_name] = viz_path
            
            # Ensure all content is properly structured
            html_content = content.copy()
            html_content['visualizations'] = viz_content
            
            rendered_html = template.render(**html_content)
            
            # Save file
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_title}_{timestamp}.html"
            file_path = self.output_dir / filename
            
            self.output_dir.mkdir(parents=True, exist_ok=True)
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(rendered_html)
            
            return file_path
            
        except Exception as e:
            logger.error(f"Error generating HTML: {str(e)}")
            return None

    def _verify_file_integrity(self, file_path: Path, format: str) -> bool:
        """Verify generated file integrity - FIXED: More lenient checks"""
        try:
            if not file_path.exists():
                logger.error(f"File does not exist: {file_path}")
                return False
            
            file_size = file_path.stat().st_size
            if file_size < 50:  # Reduced minimum size
                logger.error(f"File too small: {file_size} bytes")
                return False
            
            # Basic format-specific checks
            if format.lower() == 'pdf':
                with open(file_path, 'rb') as f:
                    header = f.read(4)
                    return header == b'%PDF'
            elif format.lower() == 'docx':
                with open(file_path, 'rb') as f:
                    header = f.read(4)
                    return header == b'PK\x03\x04'
            elif format.lower() == 'html':
                return file_size > 100  # HTML can be small but valid
            
            return True
        except Exception as e:
            logger.error(f"File integrity verification failed: {str(e)}")
            return False

    def _clean_text(self, text: str) -> str:
        """Clean text for safe generation"""
        if not text:
            return ""
        
        # Remove markdown but keep text content
        text = str(text)
        text = text.replace('**', '').replace('*', '').replace('`', '').replace('~', '')
        
        return text.strip()

    def _get_default_sections(self):
        return {
            'executive_summary': True,
            'data_overview': True,
            'eda_analysis': True,
            'ml_results': True,
            'explanations': True,
            'sql_insights': True,
            'recommendations': True
        }

    def _create_fallback_section(self, section_name: str) -> Dict[str, Any]:
        """Create fallback section when main content is empty"""
        return {
            'title': section_name.replace('_', ' ').title(),
            'content': [{
                'type': 'paragraph',
                'text': f'This section contains general information about {section_name.replace("_", " ")}. Specific analysis data is not available in the current session.'
            }]
        }

    def _create_fallback_content(self) -> Dict[str, Any]:
        """Create fallback content when no sections are available"""
        return {
            'title': 'Analysis Report',
            'content': [{
                'type': 'paragraph',
                'text': 'This report contains a summary of the data analysis performed. Please ensure all analysis steps have been completed for more detailed sections.'
            }]
        }

    def _create_error_section(self, section_name: str, error: str) -> Dict[str, Any]:
        return {
            'title': section_name.replace('_', ' ').title(),
            'content': [{
                'type': 'paragraph',
                'text': f'This section could not be generated due to an error: {error}'
            }]
        }

    # Section creation methods with proper content structure
    def _create_executive_summary(self, data: Dict[str, Any]) -> Dict[str, Any]:
        summary = {
            'title': 'Executive Summary',
            'content': []
        }
        
        try:
            if 'dataset' in data:
                df = data['dataset']
                summary['content'].append({
                    'type': 'paragraph',
                    'text': f"This analysis covers a dataset with {len(df):,} records and {len(df.columns)} features, providing comprehensive insights through exploratory data analysis and machine learning modeling."
                })
            
            if 'ml_results' in data:
                ml_data = data['ml_results']
                if 'best_model' in ml_data:
                    best_model = ml_data['best_model']
                    summary['content'].append({
                        'type': 'paragraph',
                        'text': f"The analysis identified {best_model} as the best performing model for this dataset."
                    })
            
            summary['content'].append({
                'type': 'paragraph',
                'text': "Key findings and detailed recommendations are provided in the following sections to support data-driven decision making."
            })
            
        except Exception as e:
            logger.error(f"Error creating executive summary: {str(e)}")
            summary['content'].append({
                'type': 'paragraph',
                'text': "Executive summary could not be generated due to data processing issues."
            })
        
        return summary

    def _create_data_overview(self, data: Dict[str, Any]) -> Dict[str, Any]:
        overview = {
            'title': 'Data Overview',
            'content': []
        }
        
        try:
            if 'dataset' in data:
                df = data['dataset']
                
                basic_stats = {
                    'Total Records': f"{len(df):,}",
                    'Total Features': len(df.columns),
                    'Missing Values': f"{df.isnull().sum().sum():,}",
                    'Duplicate Records': f"{df.duplicated().sum():,}",
                    'Memory Usage': f"{df.memory_usage(deep=True).sum() / 1024**2:.1f} MB"
                }
                
                overview['content'].append({
                    'type': 'table',
                    'title': 'Dataset Statistics',
                    'data': basic_stats
                })
                
                dtype_counts = df.dtypes.value_counts().to_dict()
                overview['content'].append({
                    'type': 'table',
                    'title': 'Data Types Distribution',
                    'data': {str(k): v for k, v in dtype_counts.items()}
                })
                
                # Include sample data as HTML table
                overview['content'].append({
                    'type': 'dataframe',
                    'title': 'Sample Data (First 5 Rows)',
                    'data': df.head().to_html(classes='table table-striped', index=False)
                })
            
        except Exception as e:
            logger.error(f"Error creating data overview: {str(e)}")
            overview['content'].append({
                'type': 'paragraph',
                'text': "Data overview could not be generated."
            })
        
        return overview

    def _create_eda_section(self, data: Dict[str, Any]) -> Dict[str, Any]:
        eda_section = {
            'title': 'Exploratory Data Analysis',
            'content': []
        }
        
        try:
            if 'dataset' in data:
                df = data['dataset']
                
                # Statistical summary
                numeric_df = df.select_dtypes(include=[np.number])
                if not numeric_df.empty:
                    stats_summary = numeric_df.describe()
                    eda_section['content'].append({
                        'type': 'dataframe',
                        'title': 'Statistical Summary - Numeric Features',
                        'data': stats_summary.to_html(classes='table table-striped')
                    })
                
                # Correlation insights
                if len(numeric_df.columns) > 1:
                    corr_matrix = numeric_df.corr()
                    high_corr_pairs = []
                    
                    for i in range(len(corr_matrix.columns)):
                        for j in range(i+1, len(corr_matrix.columns)):
                            corr_value = corr_matrix.iloc[i, j]
                            if abs(corr_value) > 0.7:
                                high_corr_pairs.append({
                                    'Feature 1': corr_matrix.columns[i],
                                    'Feature 2': corr_matrix.columns[j],
                                    'Correlation': f"{corr_value:.3f}"
                                })
                    
                    if high_corr_pairs:
                        eda_section['content'].append({
                            'type': 'paragraph',
                            'text': f"Found {len(high_corr_pairs)} highly correlated feature pairs (correlation > 0.7):"
                        })
                        
                        for pair in high_corr_pairs[:5]:
                            eda_section['content'].append({
                                'type': 'bullet',
                                'text': f"{pair['Feature 1']} ↔ {pair['Feature 2']}: {pair['Correlation']}"
                            })
            
        except Exception as e:
            logger.error(f"Error creating EDA section: {str(e)}")
            eda_section['content'].append({
                'type': 'paragraph',
                'text': "EDA analysis could not be generated."
            })
        
        return eda_section

    def _create_ml_section(self, data: Dict[str, Any]) -> Dict[str, Any]:
        ml_section = {
            'title': 'Machine Learning Results',
            'content': []
        }
        
        try:
            if 'ml_results' in data:
                ml_data = data['ml_results']
                
                # Best model
                if 'best_model' in ml_data:
                    best_model = ml_data['best_model']
                    ml_section['content'].append({
                        'type': 'paragraph',
                        'text': f"Best performing model: {best_model}"
                    })
                
                # Model comparison
                if 'all_models' in ml_data:
                    model_comparison = {}
                    for model_name, results in ml_data['all_models'].items():
                        metrics = results.get('test_metrics', {})
                        if 'accuracy' in metrics:
                            model_comparison[model_name] = f"{metrics['accuracy']:.3f}"
                        elif 'r2' in metrics:
                            model_comparison[model_name] = f"{metrics['r2']:.3f}"
                    
                    if model_comparison:
                        ml_section['content'].append({
                            'type': 'table',
                            'title': 'Model Performance Comparison',
                            'data': model_comparison
                        })
                
                # Feature importance
                if 'feature_importance' in ml_data:
                    importance = ml_data['feature_importance']
                    top_features = dict(list(importance.items())[:8])
                    
                    ml_section['content'].append({
                        'type': 'paragraph',
                        'text': f"Top features by importance: {', '.join(list(top_features.keys())[:3])}"
                    })
                    
                    ml_section['content'].append({
                        'type': 'table',
                        'title': 'Feature Importance (Top 8)',
                        'data': {k: f"{v:.4f}" for k, v in top_features.items()}
                    })
            
        except Exception as e:
            logger.error(f"Error creating ML section: {str(e)}")
            ml_section['content'].append({
                'type': 'paragraph',
                'text': "ML results could not be generated."
            })
        
        return ml_section

    def _create_explanations_section(self, data: Dict[str, Any]) -> Dict[str, Any]:
        explanations_section = {
            'title': 'Model Explainability',
            'content': []
        }
        
        try:
            explanations_section['content'].append({
                'type': 'paragraph',
                'text': "Model explainability provides insights into how machine learning models make decisions, enhancing transparency and trust in AI systems."
            })
            
            if 'ml_results' in data and 'feature_importance' in data['ml_results']:
                importance = data['ml_results']['feature_importance']
                top_features = list(importance.keys())[:3]
                
                explanations_section['content'].append({
                    'type': 'paragraph',
                    'text': f"The model primarily relies on these key features: {', '.join(top_features)}."
                })
            
            # Business insights
            insights = [
                "Feature importance analysis reveals the relative contribution of each variable to model predictions",
                "Understanding model behavior helps in validating results and identifying potential biases",
                "Explainable AI techniques support regulatory compliance and ethical AI practices"
            ]
            
            for insight in insights:
                explanations_section['content'].append({
                    'type': 'bullet',
                    'text': insight
                })
            
        except Exception as e:
            logger.error(f"Error creating explanations section: {str(e)}")
            explanations_section['content'].append({
                'type': 'paragraph',
                'text': "Model explanations could not be generated."
            })
        
        return explanations_section

    def _create_sql_section(self, data: Dict[str, Any]) -> Dict[str, Any]:
        sql_section = {
            'title': 'SQL Query Insights',
            'content': []
        }
        
        try:
            if 'chat_history' in data and data['chat_history']:
                sql_section['content'].append({
                    'type': 'paragraph',
                    'text': f"A total of {len(data['chat_history'])} queries were executed during the analysis session."
                })
                
                # Show recent queries
                for i, query_info in enumerate(data['chat_history'][:2]):
                    sql_section['content'].append({
                        'type': 'paragraph',
                        'text': f"Query {i+1}: {query_info.get('query', 'N/A')[:100]}..."
                    })
            else:
                sql_section['content'].append({
                    'type': 'paragraph',
                    'text': "No SQL queries were executed during this analysis session."
                })
            
        except Exception as e:
            logger.error(f"Error creating SQL section: {str(e)}")
            sql_section['content'].append({
                'type': 'paragraph',
                'text': "SQL insights could not be generated."
            })
        
        return sql_section

    def _create_recommendations_section(self, data: Dict[str, Any]) -> Dict[str, Any]:
        recommendations_section = {
            'title': 'Recommendations',
            'content': []
        }
        
        try:
            recommendations_section['content'].append({
                'type': 'paragraph',
                'text': "Based on the comprehensive analysis, the following recommendations are provided:"
            })
            
            recommendations = [
                "Implement the identified best performing model for production use",
                "Monitor model performance regularly and establish retraining procedures",
                "Consider feature engineering to improve model accuracy and robustness",
                "Establish data quality monitoring to maintain model performance over time",
                "Document model decisions and maintain version control for reproducibility"
            ]
            
            for i, recommendation in enumerate(recommendations, 1):
                recommendations_section['content'].append({
                    'type': 'numbered',
                    'number': i,
                    'text': recommendation
                })
            
        except Exception as e:
            logger.error(f"Error creating recommendations section: {str(e)}")
            recommendations_section['content'].append({
                'type': 'paragraph',
                'text': "Recommendations could not be generated."
            })
        
        return recommendations_section